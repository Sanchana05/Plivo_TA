import pdfplumber
import docx
import requests
from readability import Document
import openai
import os
from dotenv import load_dotenv
load_dotenv()

OPENAI_KEY = os.getenv("OPENAI_API_KEY")
if OPENAI_KEY:
    openai.api_key = OPENAI_KEY

def parse_pdf(path):
    text = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            txt = page.extract_text()
            if txt:
                text.append(txt)
    return "\n".join(text)

def parse_docx(path):
    doc = docx.Document(path)
    text = "\n".join([p.text for p in doc.paragraphs])
    return text

def fetch_url(url):
    r = requests.get(url, timeout=15)
    doc = Document(r.text)
    return doc.summary() or doc.content()

def parse_and_summarize(path=None, url=None):
    if url:
        text = fetch_url(url)
    else:
        if path.lower().endswith(".pdf"):
            text = parse_pdf(path)
        elif path.lower().endswith(".docx") or path.lower().endswith(".doc"):
            text = parse_docx(path)
        else:
            # try to read as plain text
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
    return summarize_text(text)

def summarize_text(text: str, max_tokens=300):
    text = text.strip()
    if not text:
        return ""
    if OPENAI_KEY:
        prompt = f"Summarize the following content concisely (3-6 sentences):\n\n{text[:20000]}"
        resp = openai.ChatCompletion.create(
            model="gpt-4o-mini",  # change as available
            messages=[{"role":"user","content":prompt}],
            max_tokens=256,
            temperature=0.2
        )
        summary = resp['choices'][0]['message']['content'].strip()
        return summary
    else:
        # fallback: very simple extractive summary - first 3 paragraphs
        paras = [p.strip() for p in text.split("\n\n") if p.strip()]
        return "\n\n".join(paras[:3])
